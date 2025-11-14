import io
import base64
try:
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None

from typing import List, Optional
from sayou.document.interfaces.base_parser import BaseDocumentParser
from sayou.document.models import (
    Document, Page, BaseElement, TextElement, TableElement, 
    ImageElement, TableCell, BoundingBox, ElementMetadata
)

class DocxParser(BaseDocumentParser):
    component_name = "DocxParser"
    SUPPORTED_TYPES = [".docx"]

    def _parse(self, file_bytes: bytes, file_name: str, **kwargs) -> Document:
        if DocxDocument is None:
            raise ImportError("python-docx is required.")

        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            raise ValueError(f"Failed to load DOCX: {e}")

        body_elements: List[BaseElement] = []
        header_elements: List[BaseElement] = []
        footer_elements: List[BaseElement] = []
        
        current_page_num = 1 
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                mixed_elements = self._process_paragraph_with_images(
                    para, doc, current_page_num, f"p{current_page_num}:body:para{id(para)}"
                )
                body_elements.extend(mixed_elements)
            
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                table_elem = self._process_table(
                    table, current_page_num, f"p{current_page_num}:body:tbl{id(table)}"
                )
                if table_elem:
                    body_elements.append(table_elem)

        try:
            for section in doc.sections:
                for para in section.header.paragraphs:
                    header_elements.extend(self._process_paragraph_with_images(
                        para, doc, current_page_num, f"p{current_page_num}:header:para{id(para)}"
                    ))
                for table in section.header.tables:
                    header_elements.append(self._process_table(
                        table, current_page_num, f"p{current_page_num}:header:tbl{id(table)}"
                    ))
                
                for para in section.footer.paragraphs:
                    footer_elements.extend(self._process_paragraph_with_images(
                        para, doc, current_page_num, f"p{current_page_num}:footer:para{id(para)}"
                    ))
                for table in section.footer.tables:
                    footer_elements.append(self._process_table(
                        table, current_page_num, f"p{current_page_num}:footer:tbl{id(table)}"
                    ))
        except Exception as e:
            self._log(f"Failed to parse headers/footers: {e}", level="warning")


        page_obj = Page(
            page_num=current_page_num,
            elements=body_elements,
            header_elements=header_elements,
            footer_elements=footer_elements,
            text="\n".join([e.text for e in body_elements if hasattr(e, "text")])
        )

        return Document(
            file_name=file_name,
            file_id=file_name,
            doc_type="word",
            page_count=1,
            pages=[page_obj]
        )

    def _process_paragraph_with_images(
        self,
        para: 'Paragraph',
        doc: 'DocxDocumentObject',
        page_num: int,
        meta_id_base: str
    ) -> List[BaseElement]:
        results = []
        current_text = ""
        style_name = para.style.name if para.style else "Normal"
        ocr_enabled = True

        for run in para.runs:
            if run.text:
                current_text += run.text

            if 'drawing' in run._element.xml:
                if current_text.strip():
                    results.append(TextElement(
                        id=f"{meta_id_base}:text{len(results)}",
                        type="text",
                        text=current_text.strip(),
                        meta=ElementMetadata(page_num=page_num, id=f"{meta_id_base}:text{len(results)}"),
                        raw_attributes={"style": style_name}
                    ))
                    current_text = "" 

                image_elem = self._extract_inline_image(run, doc, page_num, ocr_enabled)
                if image_elem:
                    image_elem.id = f"{meta_id_base}:img{len(results)}"
                    image_elem.meta.id = image_elem.id
                    results.append(image_elem)

        if current_text.strip():
            results.append(TextElement(
                id=f"{meta_id_base}:text_end",
                type="text",
                text=current_text.strip(),
                meta=ElementMetadata(page_num=page_num, id=f"{meta_id_base}:text_end"),
                raw_attributes={"style": style_name}
            ))
            
        return results

    def _extract_inline_image(self, run, doc, page_num, ocr_enabled) -> Optional[ImageElement]:
        """_process_image_data 헬퍼 사용"""
        try:
            blip_elements = run._element.findall('.//a:blip', namespaces=run._element.nsmap)
            for blip in blip_elements:
                rId = blip.get(qn('r:embed'))
                if rId:
                    image_part = doc.part.related_parts[rId]
                    image_bytes = image_part.blob
                    
                    return self._process_image_data(
                        image_bytes=image_bytes,
                        img_format=image_part.content_type.split('/')[-1] or "png", 
                        elem_id=f"img:{rId}",
                        page_num=page_num,
                        ocr_enabled=ocr_enabled
                    )
        except Exception as e:
            self._log(f"DOCX image extraction error: {e}")
            pass
        return None

    def _process_table(self, table: 'Table', page_num: int, meta_id: str) -> TableElement:
        """[수정] TableCell 객체를 생성하고 meta를 올바르게 설정"""
        rows_data = [] 
        rows_cells = [] 

        for r_idx, row in enumerate(table.rows):
            current_row_data = []
            current_row_cells = []
            
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                current_row_data.append(cell_text)
                
                # TODO: v0.1.0+ - 셀 병합(merge) 정보 파악
                cell_obj = TableCell(
                    text=cell_text,
                    row_span=1, 
                    col_span=1,
                    is_header=(r_idx == 0) # 단순 추측 (v0.1.0+ 개선 필요)
                )
                current_row_cells.append(cell_obj)
            
            rows_data.append(current_row_data)
            rows_cells.append(current_row_cells)

        meta = ElementMetadata(
            page_num=page_num,
            id=meta_id
        )

        return TableElement(
            id=meta.id,
            type="table",
            data=rows_data,
            cells=rows_cells,
            meta=meta
        )