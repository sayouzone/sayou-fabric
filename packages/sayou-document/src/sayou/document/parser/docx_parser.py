import io

try:
    from docx import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    DocxDocument = None

from typing import List, Optional

from sayou.core.registry import register_component

from ..interfaces.base_parser import BaseDocumentParser
from ..models import (
    BaseElement,
    BoundingBox,
    Document,
    ElementMetadata,
    ImageElement,
    Page,
    TableCell,
    TableElement,
    TextElement,
)


@register_component("parser")
class DocxParser(BaseDocumentParser):
    """
    (Tier 2) Parser for Microsoft Word (.docx) documents.

    Uses 'python-docx' to traverse the document tree, extracting paragraphs,
    tables, and inline images from the body, headers, and footers.
    It preserves semantic structure (Headings, Lists) in 'raw_attributes'.
    """

    component_name = "DocxParser"
    SUPPORTED_TYPES = [".docx", ".doc"]
    NAMESPACES = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    @classmethod
    def can_handle(cls, file_bytes: bytes, file_name: str) -> float:
        """
        Checks for Zip signature (PK..) or OLE signature.
        """
        # DOCX (Zip)
        if file_bytes.startswith(b"PK\x03\x04") and file_name.lower().endswith(".docx"):
            return 1.0
        # Legacy DOC
        if file_bytes.startswith(b"\xd0\xcf\x11\xe0"):
            return 1.0
        # Extension fallback
        if any(file_name.lower().endswith(t) for t in cls.SUPPORTED_TYPES):
            return 0.8
        return 0.0

    def _do_parse(self, file_bytes: bytes, file_name: str, **kwargs) -> Document:
        """
        Parse DOCX bytes into a structured Document.

        Args:
            file_bytes (bytes): The binary content of the .docx file.
            file_name (str): Original filename.
            **kwargs: Options passed to image processing (e.g., ocr_enabled).

        Returns:
            Document: A document object with 'doc_type="word"'.
        """
        if DocxDocument is None:
            raise ImportError("python-docx is required.")

        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            raise ValueError(f"Failed to load DOCX: {e}")

        body_elements: List[BaseElement] = []
        header_elements: List[BaseElement] = []
        footer_elements: List[BaseElement] = []

        current_page_num = 1

        for element in doc.element.body:
            if element.tag.endswith("p"):
                para = Paragraph(element, doc)
                mixed_elements = self._process_paragraph_with_images(
                    para,
                    doc,
                    current_page_num,
                    f"p{current_page_num}:body:para{id(para)}",
                )
                body_elements.extend(mixed_elements)

            elif element.tag.endswith("tbl"):
                table = Table(element, doc)
                table_elem = self._process_table(
                    table, current_page_num, f"p{current_page_num}:body:tbl{id(table)}"
                )
                if table_elem:
                    body_elements.append(table_elem)

        try:
            for section in doc.sections:
                for para in section.header.paragraphs:
                    header_elements.extend(
                        self._process_paragraph_with_images(
                            para,
                            doc,
                            current_page_num,
                            f"p{current_page_num}:header:para{id(para)}",
                        )
                    )
                for table in section.header.tables:
                    header_elements.append(
                        self._process_table(
                            table,
                            current_page_num,
                            f"p{current_page_num}:header:tbl{id(table)}",
                        )
                    )

                for para in section.footer.paragraphs:
                    footer_elements.extend(
                        self._process_paragraph_with_images(
                            para,
                            doc,
                            current_page_num,
                            f"p{current_page_num}:footer:para{id(para)}",
                        )
                    )
                for table in section.footer.tables:
                    footer_elements.append(
                        self._process_table(
                            table,
                            current_page_num,
                            f"p{current_page_num}:footer:tbl{id(table)}",
                        )
                    )
        except Exception as e:
            self._log(f"Failed to parse headers/footers: {e}", level="warning")

        page_obj = Page(
            page_num=current_page_num,
            elements=body_elements,
            header_elements=header_elements,
            footer_elements=footer_elements,
            text="\n".join([e.text for e in body_elements if hasattr(e, "text")]),
        )

        return Document(
            file_name=file_name,
            file_id=file_name,
            doc_type="word",
            page_count=1,
            pages=[page_obj],
        )

    def _process_paragraph_with_images(
        self,
        para: "Paragraph",
        doc: "DocxDocumentObject",
        page_num: int,
        meta_id_base: str,
    ) -> List[BaseElement]:
        """
        Extract text and inline images from a single paragraph.

        This method handles split runs (mixed text and images) and
        identifies semantic styles (Heading 1-9, List Bullet/Number).

        Args:
            para (Paragraph): The python-docx Paragraph object.
            doc (Document): The root document object (for image part access).
            page_num (int): Current page number (always 1 for flow documents).
            meta_id_base (str): Base ID string for generating element IDs.

        Returns:
            List[BaseElement]: A list of TextElement and ImageElement objects.
        """
        results = []
        current_text = ""
        style_name = para.style.name if para.style else "Normal"
        raw_attrs = {"style": style_name}
        try:
            if (
                para.style.name.startswith(("List", "목록"))
                or para._p.pPr.numPr is not None
            ):
                raw_attrs["semantic_type"] = "list"
                try:
                    raw_attrs["list_level"] = para._p.pPr.numPr.ilvl.val
                except AttributeError:
                    raw_attrs["list_level"] = 0
            elif (
                para.style.style_id.startswith("Heading")
                and para.style.type == WD_STYLE_TYPE.PARAGRAPH
            ):
                level_str = "".join(filter(str.isdigit, para.style.style_id))
                if level_str:
                    raw_attrs["semantic_type"] = "heading"
                    raw_attrs["heading_level"] = int(level_str)
        except Exception as e:
            self._log(f"Semantic style parsing error: {e}")

        ocr_enabled = True
        for run in para.runs:
            if run.text:
                current_text += run.text

            if "drawing" in run._element.xml:
                if current_text.strip():
                    results.append(
                        TextElement(
                            id=f"{meta_id_base}:text{len(results)}",
                            type="text",
                            text=current_text.strip(),
                            meta=ElementMetadata(
                                page_num=page_num,
                                id=f"{meta_id_base}:text{len(results)}",
                            ),
                            raw_attributes=raw_attrs,
                        )
                    )
                    current_text = ""

                image_elem = self._extract_inline_image(run, doc, page_num, ocr_enabled)
                if image_elem:
                    image_elem.id = f"{meta_id_base}:img{len(results)}"
                    image_elem.meta.id = image_elem.id
                    results.append(image_elem)

        if current_text.strip():
            results.append(
                TextElement(
                    id=f"{meta_id_base}:text_end",
                    type="text",
                    text=current_text.strip(),
                    meta=ElementMetadata(
                        page_num=page_num, id=f"{meta_id_base}:text_end"
                    ),
                    raw_attributes=raw_attrs,
                )
            )

        return results

    def _extract_inline_image(
        self, run, doc, page_num, ocr_enabled
    ) -> Optional[ImageElement]:
        """
        Extract an image embedded within a text run.

        Searches for 'blip' elements in the run's XML and retrieves
        the binary data from the related document part.

        Args:
            run (Run): The text run containing the image.
            doc (Document): The root document.
            page_num (int): Page number.
            ocr_enabled (bool): Whether to perform OCR on the extracted image.

        Returns:
            Optional[ImageElement]: The extracted image element, or None.
        """
        try:
            blip_elements = run._element.findall(
                ".//a:blip", namespaces=self.NAMESPACES
            )
            for blip in blip_elements:
                rId = blip.get(qn("r:embed"))
                if rId:
                    image_part = doc.part.related_parts[rId]
                    image_bytes = image_part.blob

                    return self._process_image_data(
                        image_bytes=image_bytes,
                        img_format=image_part.content_type.split("/")[-1] or "png",
                        elem_id=f"img:{rId}",
                        page_num=page_num,
                        ocr_enabled=ocr_enabled,
                    )
        except Exception as e:
            self._log(f"DOCX image extraction error: {e}")
            pass
        return None

    def _process_table(
        self, table: "Table", page_num: int, meta_id: str
    ) -> TableElement:
        """
        Convert a DOCX table into a TableElement.

        Iterates through rows and cells to build a 2D data grid and
        detailed cell objects.

        Args:
            table (Table): The python-docx Table object.
            page_num (int): Page number.
            meta_id (str): Unique ID for the table.

        Returns:
            TableElement: The structured table.
        """
        rows_data = []
        rows_cells = []

        for r_idx, row in enumerate(table.rows):
            current_row_data = []
            current_row_cells = []

            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                current_row_data.append(cell_text)

                # TODO: v0.1.0+ - 셀 병합(merge) 정보 파악
                cell_obj = TableCell(
                    text=cell_text,
                    row_span=1,
                    col_span=1,
                    is_header=(r_idx == 0),  # 단순 추측 (v0.1.0+ 개선 필요)
                )
                current_row_cells.append(cell_obj)

            rows_data.append(current_row_data)
            rows_cells.append(current_row_cells)

        meta = ElementMetadata(page_num=page_num, id=meta_id)

        return TableElement(
            id=meta.id, type="table", data=rows_data, cells=rows_cells, meta=meta
        )
